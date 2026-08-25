"""
Unit tests for document format loaders (PDF, DOCX, TXT, Markdown) and error handling.
"""

import os
import tempfile
import pytest
import docx
from pypdf import PdfWriter

from app.services.ingestion.pdf_loader import PDFLoader
from app.services.ingestion.docx_loader import DocxLoader
from app.services.ingestion.txt_loader import TxtLoader
from app.services.ingestion.markdown_loader import MarkdownLoader


@pytest.mark.asyncio
async def test_txt_loader_success():
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8") as f:
        f.write("Section Header\n\nThis is a test paragraph for TXT loader.\nSecond line of paragraph.")
        temp_path = f.name

    try:
        loader = TxtLoader()
        doc = await loader.load(temp_path, "text/plain")
        assert doc.mime_type == "text/plain"
        assert "test paragraph" in doc.raw_text
        assert len(doc.elements) == 1
    finally:
        os.remove(temp_path)


@pytest.mark.asyncio
async def test_markdown_loader_structure():
    content = """# Main Header

This is intro text.

## Section 1

- List item 1
- List item 2

### Sub Header

Detailed paragraph.
"""
    with tempfile.NamedTemporaryFile(suffix=".md", delete=False, mode="w", encoding="utf-8") as f:
        f.write(content)
        temp_path = f.name

    try:
        loader = MarkdownLoader()
        doc = await loader.load(temp_path, "text/markdown")
        assert len(doc.elements) == 6
        assert doc.elements[0].text == "Main Header"
        assert doc.elements[0].heading_level == 1
        assert doc.elements[2].text == "Section 1"
        assert doc.elements[2].heading_level == 2
    finally:
        os.remove(temp_path)


@pytest.mark.asyncio
async def test_docx_loader_structure():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        temp_path = f.name

    doc = docx.Document()
    doc.add_heading("Title Heading", level=1)
    doc.add_paragraph("First paragraph in docx.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header1"
    table.cell(0, 1).text = "Header2"
    table.cell(1, 0).text = "Val1"
    table.cell(1, 1).text = "Val2"
    doc.save(temp_path)

    try:
        loader = DocxLoader()
        extracted = await loader.load(temp_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert len(extracted.elements) >= 2
        assert extracted.elements[0].text == "Title Heading"
        assert extracted.elements[0].heading_level == 1
    finally:
        os.remove(temp_path)


@pytest.mark.asyncio
async def test_pdf_loader_structure():
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        temp_path = f.name

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with open(temp_path, "wb") as f_out:
        writer.write(f_out)

    try:
        loader = PDFLoader()
        extracted = await loader.load(temp_path, "application/pdf")
        assert extracted.metadata["total_pages"] == 1
    finally:
        os.remove(temp_path)


@pytest.mark.asyncio
async def test_corrupted_pdf_error_handling():
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False, mode="wb") as f:
        f.write(b"NOT A REAL PDF STREAM")
        temp_path = f.name

    try:
        loader = PDFLoader()
        with pytest.raises(ValueError, match="Corrupted or invalid PDF"):
            await loader.load(temp_path, "application/pdf")
    finally:
        os.remove(temp_path)
