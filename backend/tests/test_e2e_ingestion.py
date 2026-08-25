"""
End-to-end integration test verifying real format document parsing, text extraction,
normalization, structure detection, semantic chunking, and database persistence.
"""

import os
import tempfile
import pytest
import docx
from pypdf import PdfWriter

from app.db.models.document import Document
from app.db.models.workspace import Workspace
from app.repositories.document_chunk_repository import DocumentChunkRepository
from app.repositories.document_repository import DocumentRepository
from app.services.ingestion_service import IngestionService


@pytest.mark.asyncio
async def test_e2e_real_pdf_ingestion_flow(db_session, test_user):
    ws = Workspace(user_id=test_user.id, title="PDF Test WS", research_mode="Deep")
    db_session.add(ws)
    await db_session.commit()

    # Generate real multi-page PDF file
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    writer.add_blank_page(width=300, height=300)


    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        temp_path = f.name
        writer.write(f)

    try:
        doc = Document(
            workspace_id=ws.id,
            filename="multi_page.pdf",
            mime_type="application/pdf",
            file_size=os.path.getsize(temp_path),
            storage_key=temp_path,
            status="uploaded",
        )
        db_session.add(doc)
        await db_session.commit()

        doc_repo = DocumentRepository(db_session)
        chunk_repo = DocumentChunkRepository(db_session)
        service = IngestionService(doc_repo, chunk_repo)

        success = await service.ingest_document(doc.id, ws.id)
        assert success is True

        await db_session.refresh(doc)
        assert doc.status == "processed"
        assert doc.error_message is None

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


@pytest.mark.asyncio
async def test_e2e_real_docx_ingestion_flow(db_session, test_user):
    ws = Workspace(user_id=test_user.id, title="DOCX Test WS", research_mode="Deep")
    db_session.add(ws)
    await db_session.commit()

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        temp_path = f.name

    doc_file = docx.Document()
    doc_file.add_heading("Deep Research Architecture Overview", level=1)
    doc_file.add_paragraph("This document outlines the multi-agent research pipeline.")
    doc_file.add_heading("Ingestion Pipeline Component", level=2)
    doc_file.add_paragraph("Documents are converted into semantic chunks with headings and metadata.")
    doc_file.save(temp_path)

    try:
        doc = Document(
            workspace_id=ws.id,
            filename="architecture.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size=os.path.getsize(temp_path),
            storage_key=temp_path,
            status="uploaded",
        )
        db_session.add(doc)
        await db_session.commit()

        doc_repo = DocumentRepository(db_session)
        chunk_repo = DocumentChunkRepository(db_session)
        service = IngestionService(doc_repo, chunk_repo)

        success = await service.ingest_document(doc.id, ws.id)
        assert success is True

        await db_session.refresh(doc)
        assert doc.status == "processed"
        assert doc.chunk_count > 0

        chunks = await chunk_repo.get_chunks_by_document(doc.id, ws.id)
        assert len(chunks) == doc.chunk_count
        assert chunks[0].section_heading is not None
        assert "Deep Research Architecture Overview" in chunks[0].content or "Ingestion Pipeline Component" in chunks[0].content

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
