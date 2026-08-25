"""
Word (DOCX) file loader using python-docx for paragraph, heading, and table extraction.
"""

import os
from typing import List

import docx

from app.services.ingestion.base import (
    BaseLoader,
    ExtractedDocument,
    ExtractedElement,
    NodeType,
)


class DocxLoader(BaseLoader):
    async def load(self, file_path: str, mime_type: str) -> ExtractedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at path: {file_path}")

        elements: List[ExtractedElement] = []
        raw_text_parts: List[str] = []

        try:
            doc = docx.Document(file_path)

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                raw_text_parts.append(text)
                style_name = para.style.name if para.style else ""

                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.replace("Heading", "").strip())
                    except ValueError:
                        level = 1
                    elements.append(
                        ExtractedElement(
                            text=text,
                            element_type=NodeType.HEADING,
                            heading_level=level,
                            section_heading=text,
                        )
                    )
                else:
                    elements.append(
                        ExtractedElement(
                            text=text,
                            element_type=NodeType.PARAGRAPH,
                        )
                    )

            # Process tables in document
            for table in doc.tables:
                table_rows: List[str] = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        table_rows.append(" | ".join(row_cells))
                if table_rows:
                    table_text = "\n".join(table_rows)
                    raw_text_parts.append(table_text)
                    elements.append(
                        ExtractedElement(
                            text=table_text,
                            element_type=NodeType.TABLE,
                        )
                    )

            full_text = "\n\n".join(raw_text_parts)
            return ExtractedDocument(
                raw_text=full_text,
                elements=elements,
                mime_type=mime_type,
                metadata={"file_path": file_path},
            )
        except Exception as e:
            raise ValueError(f"Corrupted or invalid DOCX file: {str(e)}") from e
